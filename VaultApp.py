"""
Streamlit Web Application for Arch Thrust Line Analysis
Infrastructure-only wrapper around the solver code in Vault3.py

Usage:
  streamlit run arch_thrust_app.py
"""

import streamlit as st
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ============================================================
# Import ALL calculations from your solver module
# ============================================================
from Vault import (
    LoadGrid,
    beam_MV,
    find_best_H,
    contact_clusters,
    plot_results,
    plot_loading,
    solve_from_intrados,   # your fixed version (with remesh_loadgrid + relax)
)

# ------------------------------------------------------------
# Small UI helper
# ------------------------------------------------------------
def editor_to_float_array(ed: pd.DataFrame, ncol: int) -> np.ndarray:
    """
    Convert a Streamlit data_editor dataframe into a clean float ndarray
    with shape (n, ncol). Drops rows with any non-numeric/NaN entries.
    """
    if ed is None or len(ed) == 0:
        return np.array([]).reshape(0, ncol)

    ed = ed.copy().iloc[:, :ncol]

    def scalar(v):
        if isinstance(v, (list, tuple, np.ndarray)):
            return v[0] if len(v) else np.nan
        return v

    for c in ed.columns:
        ed[c] = ed[c].map(scalar)
        ed[c] = pd.to_numeric(ed[c], errors="coerce")

    ed = ed.dropna(how="any")
    return ed.to_numpy(dtype=float) if len(ed) else np.array([]).reshape(0, ncol)


# ============================================================
# DOCX REPORT (kept here; no numeric solving done inside)
# ============================================================
def build_docx_python(
    *,
    L: float, f: float, dx: float,
    sol: dict, clusters: list[tuple[int, int, int]],
    udls: np.ndarray, vdls: np.ndarray, pls: np.ndarray,
    RA: float, RB: float,
    info: dict,
    load_png, res_png,
    beam_x: np.ndarray, beam_V: np.ndarray, beam_M: np.ndarray, beam_RA: float, beam_RB: float,
    b_eff: float, fcd: float,
    N_max: float, a_max: float, x_a: float, t_total: float
) -> bytes:
    """
    Create a Word report from already-computed results.
    (No analysis is performed here; it only formats inputs.)
    """
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    style.font.size = Pt(10)

    def add_title(text: str):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(18)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_meta_line(text: str):
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_h(level: int, text: str):
        doc.add_heading(text, level=level)

    def add_kv_table(rows: list[tuple[str, str]]):
        t = doc.add_table(rows=len(rows), cols=2)
        t.style = "Table Grid"
        for i, (k, v) in enumerate(rows):
            t.cell(i, 0).text = str(k)
            t.cell(i, 1).text = str(v)
        return t

    def add_table(headers: list[str], data_rows: list[list[str]]):
        t = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
        t.style = "Table Grid"
        for j, h in enumerate(headers):
            t.cell(0, j).text = h
        for i, row in enumerate(data_rows, start=1):
            for j, val in enumerate(row):
                t.cell(i, j).text = val
        return t

    # Basic checks purely from passed arrays
    M0 = float(beam_M[0]) if len(beam_M) else float("nan")
    ML = float(beam_M[-1]) if len(beam_M) else float("nan")

    e = np.asarray(sol["e"], dtype=float)
    delta = float(sol["delta"])
    emax = float(np.max(np.abs(e))) if len(e) else float("nan")

    add_title("Arch Thrust Line Analysis (Minimax / Lower-Bound)")
    add_meta_line(f"{info.get('project','')}  |  {info.get('author','')}  |  {info.get('date','')}")
    doc.add_paragraph("")

    add_h(1, "1. Geometry")
    add_kv_table([
        ("Span, L", f"{L:.3f} m"),
        ("Rise, f", f"{f:.3f} m"),
        ("Rise/Span", f"{(f/L):.3%}"),
        ("Discretisation dx (loading)", f"{dx:.4f} m"),
        ("Thrust evaluation points", f"{len(sol['x'])}"),
    ])

    add_h(1, "2. Loading input")
    add_h(2, "2.1 UDLs")
    if len(udls):
        add_table(["w", "x0", "x1"], [[f"{r[0]:.4g}", f"{r[1]:.4g}", f"{r[2]:.4g}"] for r in udls])
    else:
        doc.add_paragraph("None")

    add_h(2, "2.2 VDLs")
    if len(vdls):
        add_table(["w0", "w1", "x0", "x1"], [[f"{r[0]:.4g}", f"{r[1]:.4g}", f"{r[2]:.4g}", f"{r[3]:.4g}"] for r in vdls])
    else:
        doc.add_paragraph("None")

    add_h(2, "2.3 Point loads")
    if len(pls):
        add_table(["P", "x"], [[f"{r[0]:.4g}", f"{r[1]:.4g}"] for r in pls])
    else:
        doc.add_paragraph("None")

    add_h(2, "2.4 Loading diagram")
    doc.add_picture(str(load_png), width=Inches(6.5))

    add_h(1, "3. Equivalent beam reactions")
    add_kv_table([
        ("R_A", f"{RA:.3f} kN"),
        ("R_B", f"{RB:.3f} kN"),
        ("R_A + R_B", f"{(RA+RB):.3f} kN"),
    ])

    add_h(1, "4. Results")
    add_kv_table([
        ("H*", f"{sol['H']:.3f} kN"),
        ("a", f"{sol['a']:.8g}"),
        ("b", f"{sol['b']:.8g}"),
        ("δ", f"{sol['delta']*1000:.3f} mm"),
        ("t_req = 2δ", f"{sol['t_req']*1000:.3f} mm"),
    ])

    add_h(2, "4.1 Compression strut sizing (capacity check)")
    add_kv_table([
        ("Effective width b", f"{b_eff:.3f} m"),
        ("Compression capacity f_cd", f"{fcd:.3f} MPa"),
        ("N_max", f"{N_max:.3f} kN"),
        ("a_max", f"{a_max*1000:.1f} mm"),
        ("at x", f"{x_a:.3f} m"),
        ("t_total = t_req + a_max", f"{t_total*1000:.1f} mm"),
    ])

    add_h(2, "4.2 Contact (hinge) locations")
    if clusters:
        xx = np.asarray(sol["x"], dtype=float)
        rows = []
        for k, (i0, i1, sgn) in enumerate(clusters, start=1):
            xh = 0.5 * (xx[i0] + xx[i1])
            side = "Extrados" if sgn > 0 else "Intrados"
            rows.append([str(k), f"{xh:.4f}", side])
        add_table(["#", "x (m)", "Side"], rows)
    else:
        doc.add_paragraph("No contact clusters detected (check tolerances / discretisation).")

    add_h(1, "5. Numerical sanity")
    add_kv_table([
        ("Beam end moment M(0)", f"{M0:.6g}"),
        ("Beam end moment M(L)", f"{ML:.6g}"),
        ("max|e(x)|", f"{emax*1000:.3f} mm"),
        ("δ", f"{delta*1000:.3f} mm"),
        ("max|e| − δ", f"{(emax-delta)*1000:.3f} mm"),
    ])

    add_h(1, "6. Diagrams")
    doc.add_picture(str(res_png), width=Inches(6.5))

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


# ============================================================
# APP STATE + UI
# ============================================================
st.set_page_config(page_title="Arch Thrust Line", page_icon="🌉", layout="wide")

for key, val in [
    ("udls", np.array([[10.0, 0.0, 20.0]])),
    ("vdls", np.array([]).reshape(0, 4)),
    ("pls",  np.array([[50.0, 5.0], [30.0, 15.0]])),
    ("results", None),
]:
    if key not in st.session_state:
        st.session_state[key] = val

if "docx_bytes" not in st.session_state:
    st.session_state.docx_bytes = None


# ============================================================
# SIDEBAR
# ============================================================
st.title("🌉 Arch Thrust Line Analysis")

with st.sidebar:
    st.header("📋 Project")
    proj = st.text_input("Project", "Arch Analysis")
    auth = st.text_input("Author", "")
    date = st.date_input("Date")

    st.divider()
    st.header("📐 Geometry")
    inside_L = st.number_input("Intrados span L (m)", 1.0, 200.0, 25.0)
    inside_f = st.number_input("Intrados rise f (m)", 0.05, 100.0, 5.0)
    st.markdown(f"**Rise/Span:** {inside_f/inside_L:.1%}")
    dx = st.number_input("Grid spacing dx (m)", 0.005, 0.5, 0.05, format="%.3f")

    st.divider()
    st.header("🧱 Material / Capacity")
    b_eff = st.number_input("Effective width b_eff (m)", min_value=0.01, value=1.0, step=0.1)
    fcd  = st.number_input("Compression capacity f_cd (MPa)", min_value=0.0, value=0.5, step=0.1)

    st.divider()
    st.header("⚙️ Solver controls")
    H_hi = st.number_input("Max H (kN)", 100.0, 1e7, 5e4, format="%.0f")
    relax = st.slider("Relaxation (intrados iteration)", min_value=0.1, max_value=1.0, value=0.5, step=0.05)


# ============================================================
# LOAD INPUT
# ============================================================
st.header("📊 Loading")

c1, c2, c3 = st.columns(3)
with c1:
    st.subheader("UDL")
    df = pd.DataFrame(st.session_state.udls, columns=["w", "x0", "x1"])
    ed = st.data_editor(df, num_rows="dynamic", key="u", height=180)
    st.session_state.udls = editor_to_float_array(ed, 3)

with c2:
    st.subheader("VDL")
    df = pd.DataFrame(st.session_state.vdls, columns=["w0", "w1", "x0", "x1"])
    ed = st.data_editor(df, num_rows="dynamic", key="v", height=180)
    st.session_state.vdls = editor_to_float_array(ed, 4)

with c3:
    st.subheader("Point")
    df = pd.DataFrame(st.session_state.pls, columns=["P", "x"])
    ed = st.data_editor(df, num_rows="dynamic", key="p", height=180)
    st.session_state.pls = editor_to_float_array(ed, 2)

# Build discretised loads on intrados span
g = LoadGrid(inside_L, dx)

for r in st.session_state.udls:
    try:
        w, x0, x1 = map(float, r[:3])
        g.udl(w, x0, x1)
    except Exception:
        pass

for r in st.session_state.vdls:
    try:
        w0, w1, x0, x1 = map(float, r[:4])
        xa, xb = (x0, x1) if x0 <= x1 else (x1, x0)
        if xb > xa:
            g.vdl(w0, w1, xa, xb)
    except Exception:
        pass

for r in st.session_state.pls:
    try:
        P, x0 = map(float, r[:2])
        g.pl(P, x0)
    except Exception:
        pass

st.subheader("📈 Preview")
fig_l = plot_loading(g, inside_L)
st.pyplot(fig_l)
plt.close(fig_l)

try:
    trapz = np.trapezoid
except AttributeError:
    trapz = np.trapz

tot = float(trapz(g.q, g.x) + sum(P for P, _ in g.pls))
st.info(f"**Total vertical load: {tot:.1f} kN**")


# ============================================================
# RUN
# ============================================================
st.divider()

if st.button("🔍 Run Analysis", type="primary"):
    with st.spinner("Computing..."):
        try:
            # Solve using your Vault3 infrastructure (intrados -> centreline iteration)
            result = solve_from_intrados(
                inside_L,
                inside_f,
                g,
                b_eff,
                fcd,
                H_hi=H_hi,
                relax=relax,
                verbose=False,
            )

            st.session_state.results = result
            st.session_state.g_stored = g

            st.success("Done!")
        except Exception as e:
            st.session_state.results = None
            st.error(str(e))


# ============================================================
# RESULTS
# ============================================================
if st.session_state.results:
    result = st.session_state.results

    sol = result["sol"]
    L = float(result["L"])
    f = float(result["f"])
    t_geo = float(result["t_geo"])
    a_max = float(result["a_max"])
    t_total = float(result["t_total"])
    V_sol = result["V_sol"]
    N_sol = result["N_sol"]
    a_sol = result["a_sol"]
    RA = float(result["RA"])
    RB = float(result["RB"])

    clusters = contact_clusters(sol["x"], sol["e"], sol["delta"])

    st.header("📊 Results")

    cols = st.columns(5)
    cols[0].metric("H*", f"{sol['H']:.1f} kN")
    cols[1].metric("t_req (geom)", f"{t_geo*1000:.1f} mm")
    cols[2].metric("a_max (strut)", f"{a_max*1000:.1f} mm")
    cols[3].metric("t_total", f"{t_total*1000:.1f} mm")
    cols[4].metric("R_A / R_B", f"{RA:.1f} / {RB:.1f} kN")

    # Hinges / contact
    st.subheader("🔗 Contact (hinge) clusters")
    if clusters:
        xx = np.asarray(sol["x"], dtype=float)
        st.table(pd.DataFrame([
            {"#": k,
             "x": f"{0.5*(xx[i0]+xx[i1]):.3f} m",
             "Side": "EXTRADOS" if s > 0 else "INTRADOS"}
            for k, (i0, i1, s) in enumerate(clusters, 1)
        ]))
    else:
        st.info("No contact clusters detected (check tolerances / discretisation).")

    # Compressive
    i = int(np.argmax(a_sol)) if a_sol is not None else 0
    x_a = float(sol["x"][i]) if a_sol is not None else float("nan")
    N_max = float(np.max(N_sol)) if N_sol is not None else float("nan")

    cS1, cS2, cS3 = st.columns(3)
    cS1.metric("N_max", f"{N_max:.1f} kN")
    cS2.metric("a_max location", f"{x_a:.2f} m")
    cS3.metric("Rise/Span (centreline)", f"{f/L:.2%}")

    # Plots
    st.subheader("📈 Diagrams")
    fig = plot_results(sol, L, f, clusters, a_sol=a_sol, V_sol=V_sol, N_sol=N_sol)
    st.pyplot(fig)

    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    buf.seek(0)
    st.download_button("📥 Download PNG", buf, "arch.png", "image/png")
    plt.close(fig)

    st.divider()

    # DOCX export
    if st.button("📝 Print Calculations (Word)", key="make_docx"):
        with st.spinner("Generating..."):
            tmpdir = Path(tempfile.mkdtemp(prefix="arch_thrust_"))
            load_png = tmpdir / "load.png"
            res_png  = tmpdir / "res.png"

            fig_l = plot_loading(st.session_state.g_stored, inside_L)
            fig_l.savefig(load_png, dpi=200, bbox_inches="tight")
            plt.close(fig_l)

            fig_r = plot_results(sol, L, f, clusters, a_sol=a_sol, V_sol=V_sol, N_sol=N_sol)
            fig_r.savefig(res_png, dpi=200, bbox_inches="tight")
            plt.close(fig_r)

            # Beam arrays for report sanity fields
            x_b, V_b, M_b, RA_b, RB_b = beam_MV(st.session_state.g_stored)

            st.session_state.docx_bytes = build_docx_python(
                L=L, f=f, dx=dx,
                sol=sol, clusters=clusters,
                udls=st.session_state.udls, vdls=st.session_state.vdls, pls=st.session_state.pls,
                RA=RA, RB=RB,
                info={"project": proj, "author": auth, "date": str(date)},
                load_png=load_png, res_png=res_png,
                beam_x=x_b, beam_V=V_b, beam_M=M_b, beam_RA=RA_b, beam_RB=RB_b,
                b_eff=b_eff, fcd=fcd,
                N_max=N_max, a_max=a_max, x_a=x_a, t_total=t_total,
            )

    if st.session_state.docx_bytes is not None:
        st.download_button(
            "📥 Download Word",
            data=st.session_state.docx_bytes,
            file_name=f"Arch_{proj.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_docx",
        )

st.divider()
st.caption("Arch Thrust Line | Minimax | Lower Bound")